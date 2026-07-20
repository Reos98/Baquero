import os
from docx import Document
from docx.shared import Pt, Inches
from pygments import highlight
from pygments.lexers import get_lexer_by_name
from pygments.formatters import ImageFormatter

def create_code_image(code, lang, output_path):
    lexer = get_lexer_by_name(lang)
    formatter = ImageFormatter(font_size=16, line_numbers=True, style='monokai')
    with open(output_path, "wb") as f:
        f.write(highlight(code, lexer, formatter))

def create_report():
    doc = Document()
    
    # Title
    doc.add_heading('Informe de Avance: Trabajo Práctico Experimental', 0)
    
    # Intro
    doc.add_paragraph('El objetivo de este informe es documentar formalmente el progreso de la aplicación web / móvil desarrollada. En este documento se detallan las decisiones arquitectónicas, el diseño de la interfaz y la integración con el backend. A continuación, se adjuntan las descripciones de las implementaciones por cada módulo, acompañadas de las respectivas capturas de pantalla del código.')
    
    base_dir = r"C:\Users\h\Music\APPS MOVILES-20260712T212108Z-2-001\APPS MOVILES"
    
    files_to_include = [
        {
            'title': '1. Frontend - Interfaz de Usuario y Autenticación',
            'desc': 'En el módulo principal de la aplicación, desarrollado en Flutter, se ha implementado el archivo main.dart, el cual actúa como punto de entrada de la interfaz gráfica.\n\nCaracterísticas implementadas:\n- Diseño de pantallas de inicio de sesión (LoginPage) y registro (RegisterPage).\n- Integración de validación de campos para garantizar la integridad de los datos de entrada (como correo electrónico y contraseñas).\n- Interfaz amigable para la gestión de notas en la pantalla principal (HomePage) con un esquema de colores dinámico que mejora la experiencia de usuario.\n- Uso de Widgets con estado (StatefulWidgets) para gestionar de forma asincrónica la autenticación con la base de datos local y evitar el bloqueo de la interfaz mientras se procesan los datos.',
            'path': os.path.join(base_dir, 'lib-20260707T200135Z-3-001', 'lib', 'main.dart'),
            'lang': 'dart',
            'lines': slice(39, 79)
        },
        {
            'title': '2. Frontend - Persistencia de Datos Local (SQLite)',
            'desc': 'Para asegurar el funcionamiento de la aplicación en modo offline y mantener persistencia de las credenciales de sesión y notas de forma rápida, se implementó el archivo db_helper.dart.\n\nAspectos técnicos destacados:\n- Patrón Singleton: Se utiliza para mantener una única conexión con la base de datos a lo largo de todo el ciclo de vida de la aplicación.\n- Consultas SQL: Se implementaron métodos asincrónicos para operaciones CRUD (Create, Read, Update, Delete) enfocadas en las tablas de usuarios y notas.\n- Este módulo actúa como intermediario local para agilizar las cargas mientras se desarrolla la sincronización total con la nube en la siguiente fase.',
            'path': os.path.join(base_dir, 'lib-20260707T200135Z-3-001', 'lib', 'db_helper.dart'),
            'lang': 'dart',
            'lines': slice(15, 45)
        },
        {
            'title': '3. Backend - API REST (Python / Flask)',
            'desc': 'El backend ha sido construido utilizando Python con el framework Flask, ubicado en app.py, lo que nos permite un gran rendimiento y escalabilidad para la lógica del servidor.\n\nImplementaciones clave:\n- JWT (JSON Web Tokens): Implementado para manejar la seguridad en cada endpoint, requiriendo autenticación para cualquier transacción sensible y utilizando tokens de refresco.\n- Base de Datos Relacional: Migraciones automáticas mediante sqlite3 para estructurar las relaciones de Uno a Muchos entre los usuarios y sus notas.\n- Middleware de Permisos: Se incorporaron decoradores de Flask personalizados (@role_required y @permission_required) para controlar de forma granular el nivel de acceso según si el usuario es "user" o "admin".\n- Swagger UI: Documentación automatizada de los endpoints expuestos.',
            'path': os.path.join(base_dir, 'backend-20260707T200134Z-3-001', 'backend', 'app.py'),
            'lang': 'python',
            'lines': slice(326, 359)
        }
    ]
    
    for i, item in enumerate(files_to_include):
        doc.add_heading(item['title'], level=1)
        doc.add_paragraph(item['desc'])
        
        doc.add_heading('Captura de pantalla de la implementación:', level=2)
        
        try:
            with open(item['path'], 'r', encoding='utf-8') as f:
                content_lines = f.readlines()
                # Extraemos el fragmento a mostrar en la captura
                snippet = "".join(content_lines[item['lines']])
                
                # Generamos la imagen
                img_path = os.path.join(base_dir, f'snippet_{i}.png')
                create_code_image(snippet, item['lang'], img_path)
                
                # Insertamos la imagen en Word
                doc.add_picture(img_path, width=Inches(6.0))
                
                # Opcional: borrar la imagen temporal
                # os.remove(img_path)
        except Exception as e:
            doc.add_paragraph(f"[No se pudo generar la captura visual. Error: {e}]")
            
    # Conclusión
    doc.add_heading('Conclusiones y Próximos Pasos', level=1)
    doc.add_paragraph('Hasta este momento, la integración de la base de datos, el flujo de autenticación local e infraestructura del servidor backend están finalizadas con éxito. Los próximos pasos consisten en integrar el cliente móvil (Frontend) para que consuma de forma directa los endpoints del backend, permitiendo la sincronización de las notas en la nube, y realizar las pruebas finales E2E (End to End).')

    # Save the document
    output_path = os.path.join(base_dir, 'informe_de_avance_v2.docx')
    doc.save(output_path)
    print(f"Informe actualizado guardado en {output_path}")

if __name__ == '__main__':
    create_report()
