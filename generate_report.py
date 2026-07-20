import os
from docx import Document
from docx.shared import Pt, Inches

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Informe de Avance: Trabajo Práctico Experimental', 0)
    
    # Intro
    doc.add_paragraph('El objetivo de este informe es monitorear el progreso del trabajo y detallar las implementaciones realizadas hasta el momento en la aplicación web / móvil.')
    
    # Sections
    base_dir = r"C:\Users\h\Music\APPS MOVILES-20260712T212108Z-2-001\APPS MOVILES"
    
    files_to_include = [
        {
            'title': 'Frontend - Interfaz de Usuario (main.dart)',
            'path': os.path.join(base_dir, 'lib-20260707T200135Z-3-001', 'lib', 'main.dart'),
            'desc': 'Se implementó la pantalla de inicio de sesión (Login), el registro de usuarios con selección de roles, y la pantalla principal (HomePage) para gestionar notas con colores dinámicos.'
        },
        {
            'title': 'Frontend - Base de Datos Local (db_helper.dart)',
            'path': os.path.join(base_dir, 'lib-20260707T200135Z-3-001', 'lib', 'db_helper.dart'),
            'desc': 'Se configuró el Helper de SQLite para manejar las operaciones CRUD locales para usuarios y notas, permitiendo la persistencia de datos en el dispositivo.'
        },
        {
            'title': 'Backend - API REST (app.py)',
            'path': os.path.join(base_dir, 'backend-20260707T200134Z-3-001', 'backend', 'app.py'),
            'desc': 'Se desarrolló la estructura principal del backend utilizando un framework de Python para exponer los endpoints necesarios de la aplicación (creación de notas, autenticación, etc.).'
        }
    ]
    
    for item in files_to_include:
        doc.add_heading(item['title'], level=1)
        doc.add_paragraph(item['desc'])
        
        doc.add_heading('Código Implementado (Captura/Snippet):', level=2)
        
        try:
            with open(item['path'], 'r', encoding='utf-8') as f:
                content = f.read()
                # Limit content to first 50 lines to simulate a screenshot/snippet
                lines = content.split('\n')[:50]
                snippet = '\n'.join(lines) + ('\n... (código truncado)' if len(content.split('\n')) > 50 else '')
                
                p = doc.add_paragraph(snippet)
                p.style.font.name = 'Courier New'
                p.style.font.size = Pt(8)
        except Exception as e:
            doc.add_paragraph(f"Error al leer el archivo: {e}")
            
    # Save the document
    output_path = os.path.join(base_dir, 'informe_de_avance.docx')
    doc.save(output_path)
    print(f"Informe guardado en {output_path}")

if __name__ == '__main__':
    create_report()
